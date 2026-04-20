#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GenerateAgent - 生成最终 CNVD Word 报告。

说明：
1) 优先按样例模板填充内容；
2) 生成后自动把真实漏洞代码文件以 OLE 对象插入“存在漏洞的代码文件”段落；
3) OLE 若在当前会话失败，会自动尝试通过计划任务在交互会话补插。
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
import tempfile
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from agents.base_agent import BaseAgent
from core.config import config
from core.state import TaskState


class GenerateAgent(BaseAgent):
    """报告生成 Agent"""

    def __init__(self, llm_client=None):
        super().__init__("GenerateAgent", llm_client)

    def _execute(self, state: TaskState) -> Dict[str, Any]:
        report_name = state.report_name

        parsed_data = self._load_previous_output(state, "parse", "parsed.json") or {}
        github_result = self._load_previous_output(state, "github", "github_result.json") or {}
        deployment_data = self._load_previous_output(state, "deploy", "deployment.json") or {}
        cvss_result = self._load_previous_output(state, "cvss", "cvss_result.json") or {}
        sqlmap_result = self._load_previous_output(state, "sqlmap", "sqlmap_result.json") or {}

        self.logger.info(f"开始生成 CNVD 报告: {report_name}")
        trusted_report_path = self._resolve_trusted_report_path(state, report_name)

        data = self._prepare_report_data(
            report_name=report_name,
            trusted_report_path=trusted_report_path,
            parsed_data=parsed_data,
            github_result=github_result,
            deployment_data=deployment_data,
            cvss_result=cvss_result,
            sqlmap_result=sqlmap_result,
        )
        data["sqlmap_screenshot"] = self._resolve_sqlmap_screenshot_path(
            report_name=report_name,
            sqlmap_result=sqlmap_result,
            raw_path=str(data.get("sqlmap_screenshot", "") or ""),
        )

        output_path = Path(f"workspace/output/{report_name}.docx")
        output_path.parent.mkdir(parents=True, exist_ok=True)

        template_path = self._select_template_path(state, report_name)
        if template_path:
            self.logger.info(f"使用模板生成报告: {template_path}")
            shutil.copy2(template_path, output_path)
            try:
                self._fill_report_from_template(output_path, data)
            except Exception as error:
                self.logger.warning(f"模板填充失败，回退基础报告: {error}")
                self._create_basic_report(output_path, data)
        else:
            self.logger.warning("未找到可用模板，使用基础报告结构生成")
            self._create_basic_report(output_path, data)

        # 强制尝试 OLE 嵌入
        self._try_embed_vulnerable_files_as_ole(
            doc_path=output_path,
            file_objects=data.get("vulnerable_file_objects", []),
        )

        self.logger.info(f"CNVD 报告已生成: {output_path}")
        return {"output_path": str(output_path), "data": {"report_path": str(output_path)}}

    # ------------------------------ data prep ------------------------------
    def _prepare_report_data(
        self,
        report_name: str,
        trusted_report_path: Optional[Path],
        parsed_data: Dict[str, Any],
        github_result: Dict[str, Any],
        deployment_data: Dict[str, Any],
        cvss_result: Dict[str, Any],
        sqlmap_result: Dict[str, Any],
    ) -> Dict[str, Any]:
        parsed_files = parsed_data.get("vulnerable_files", [])
        deployment_files = deployment_data.get("files_to_embed", [])
        downloaded_path = str(deployment_data.get("downloaded_path", "") or "")

        vulnerable_file_objects = self._build_vulnerable_file_objects(
            parsed_files=parsed_files,
            deployment_files=deployment_files,
            downloaded_path=downloaded_path,
        )
        vulnerable_file_lines = self._build_vulnerable_file_lines(parsed_files, vulnerable_file_objects)

        source_url = self._pick_source_download_url(parsed_data, github_result)
        repo_url = str(github_result.get("repository_url", "") or "").strip()

        return {
            "product_description": str(parsed_data.get("product_description", "") or ""),
            "source_download_url": source_url,
            "repository_url": repo_url,
            "vulnerable_files": vulnerable_file_lines,
            "vulnerable_file_objects": vulnerable_file_objects,
            "cvss_score": cvss_result.get("base_score", "N/A") if cvss_result else "N/A",
            "vulnerability_principle": str(parsed_data.get("vulnerability_principle", "") or ""),
            "vulnerable_url": str(parsed_data.get("vulnerable_url", "") or ""),
            "reproduction_steps": parsed_data.get("reproduction_steps", []) or [],
            "sqlmap_command_hint": str(parsed_data.get("sqlmap_command", "") or "").strip(),
            "sqlmap_command_executed": str(sqlmap_result.get("command", "") or "").strip(),
            "sqlmap_screenshot": str(sqlmap_result.get("screenshot_path", "") or ""),
            "sqlmap_executed": bool(sqlmap_result.get("executed", False)) if sqlmap_result else False,
            "sqlmap_failure_reason": str(sqlmap_result.get("failure_reason", "") or ""),
            "sqlmap_error": str(sqlmap_result.get("error", "") or ""),
            "trusted_section_images": self._extract_section_images_from_doc(
                trusted_report_path, report_name
            )
            if trusted_report_path
            else {},
        }

    def _resolve_trusted_report_path(self, state: TaskState, report_name: str) -> Optional[Path]:
        candidates: List[Path] = []
        candidates.append(Path(str(state.input_file)))
        candidates.append(Path("可信代码库报告") / f"{report_name}.docx")
        trusted_dir = Path("可信代码库报告")
        if trusted_dir.exists():
            candidates.extend(sorted(trusted_dir.glob("*.docx")))

        seen = set()
        for candidate in candidates:
            key = str(candidate).lower()
            if key in seen:
                continue
            seen.add(key)
            try:
                if candidate.exists() and candidate.is_file():
                    return candidate.resolve()
            except Exception:
                continue
        return None

    def _extract_section_images_from_doc(
        self,
        doc_path: Optional[Path],
        report_name: str,
    ) -> Dict[str, List[str]]:
        if not doc_path:
            return {}
        try:
            from docx import Document
        except Exception as error:
            self.logger.warning(f"解析可信报告图片失败（python-docx不可用）: {error}")
            return {}

        try:
            doc = Document(doc_path)
        except Exception as error:
            self.logger.warning(f"打开可信报告失败，无法提取图片: {doc_path} | {error}")
            return {}

        safe_report_name = "".join(c for c in str(report_name) if c.isalnum() or c in ("-", "_")) or "task"
        output_dir = Path("workspace/tmp") / safe_report_name / "trusted_images"
        output_dir.mkdir(parents=True, exist_ok=True)

        section_images: Dict[str, List[str]] = {}
        current_heading = ""

        for para in doc.paragraphs:
            style_name = str(getattr(para.style, "name", "") or "")
            text = str(para.text or "").strip()
            if style_name.startswith("Heading") and text:
                current_heading = text
                continue

            target_section = self._map_source_heading_to_target(current_heading)
            if not target_section or target_section == "复现流程":
                continue

            rel_ids = re.findall(r'r:embed="([^"]+)"', para._p.xml)
            if not rel_ids:
                continue

            for rid in rel_ids:
                try:
                    image_part = doc.part.related_parts[rid]
                    blob = image_part.blob
                except Exception:
                    continue

                ext = Path(str(getattr(image_part, "partname", ""))).suffix or ".png"
                idx = len(section_images.get(target_section, [])) + 1
                image_path = output_dir / f"{target_section}_{idx}{ext}"
                try:
                    image_path.write_bytes(blob)
                    section_images.setdefault(target_section, []).append(str(image_path.resolve()))
                except Exception:
                    continue

        return section_images

    def _map_source_heading_to_target(self, heading: str) -> str:
        key = self._normalize_text_key(heading)
        if not key:
            return ""
        if "产品介绍" in key:
            return "产品介绍"
        # “代码仓库 / GitHub”章节截图不迁移到 CNVD 报告
        if "代码仓库" in key or "源码" in key or "github" in key.lower():
            return ""
        if "存在漏洞的代码文件" in key or "存在漏洞的代码" in key:
            return "存在漏洞的代码文件"
        if "漏洞评分" in key or "漏洞等级" in key or "cvss" in key.lower():
            return "漏洞评分"
        if "漏洞原理" in key:
            return "漏洞原理"
        if "漏洞url" in key:
            return "漏洞URL"
        if "复现流程" in key:
            return "复现流程"
        return ""

    def _pick_source_download_url(
        self,
        parsed_data: Dict[str, Any],
        github_result: Dict[str, Any],
    ) -> str:
        direct = str(parsed_data.get("source_download_url", "") or "").strip()
        if direct:
            return direct

        raw_text = str(parsed_data.get("raw_text", "") or "")
        match = re.search(r"https?://github\.com/[^\s]+?\.zip", raw_text)
        if match:
            return match.group(0).strip()

        return str(github_result.get("download_url", "") or "").strip()

    def _resolve_sqlmap_screenshot_path(
        self,
        report_name: str,
        sqlmap_result: Dict[str, Any],
        raw_path: str,
    ) -> str:
        candidates: List[Path] = []
        if raw_path:
            candidates.append(Path(raw_path))

        try:
            candidates.append(self.state_manager.get_step_output_path(report_name, "sqlmap", "sqlmap_result.png"))
        except Exception:
            pass

        candidates.append(Path(f"workspace/{report_name}/05_sqlmap/sqlmap_result.png"))
        candidates.append(Path("workspace") / report_name / "05_sqlmap" / "sqlmap_result.png")

        for candidate in candidates:
            try:
                if candidate and candidate.exists():
                    return str(candidate.resolve())
            except Exception:
                continue

        alt_raw = str(sqlmap_result.get("output_log", "") or "")
        if alt_raw:
            try:
                alt_dir = Path(alt_raw).parent
                alt_candidate = alt_dir / "sqlmap_result.png"
                if alt_candidate.exists():
                    return str(alt_candidate.resolve())
            except Exception:
                pass

        return ""

    def _build_vulnerable_file_objects(
        self,
        parsed_files: Any,
        deployment_files: Any,
        downloaded_path: str,
    ) -> List[Dict[str, str]]:
        items: List[Dict[str, str]] = []
        seen = set()

        if isinstance(deployment_files, list):
            for obj in deployment_files:
                if not isinstance(obj, dict):
                    continue
                source_path = str(obj.get("source_path", "") or "").strip()
                if not source_path:
                    continue
                path = Path(source_path)
                if not path.exists():
                    continue
                key = str(path.resolve()).lower()
                if key in seen:
                    continue
                seen.add(key)
                items.append(
                    {
                        "source_path": str(path.resolve()),
                        "display_name": str(obj.get("display_name", "") or path.name),
                        "original_path": str(obj.get("original_path", "") or ""),
                    }
                )

        source_root = Path(downloaded_path) if downloaded_path else None
        if isinstance(parsed_files, list) and source_root and source_root.exists():
            for raw in parsed_files:
                rel = str(raw or "").strip().lstrip("./").lstrip("\\/")
                if not rel:
                    continue
                candidate = (source_root / rel).resolve()
                if not candidate.exists():
                    continue
                key = str(candidate).lower()
                if key in seen:
                    continue
                seen.add(key)
                items.append(
                    {
                        "source_path": str(candidate),
                        "display_name": candidate.name,
                        "original_path": str(raw or ""),
                    }
                )
        return items

    def _build_vulnerable_file_lines(
        self,
        parsed_files: Any,
        file_objects: List[Dict[str, str]],
    ) -> List[str]:
        lines: List[str] = []
        seen = set()

        if isinstance(parsed_files, list):
            for item in parsed_files:
                line = str(item or "").strip()
                if not line:
                    continue
                key = line.lower()
                if key in seen:
                    continue
                seen.add(key)
                lines.append(line)

        for obj in file_objects:
            preferred = str(obj.get("original_path", "") or "").strip() or str(obj.get("display_name", "") or "").strip()
            if not preferred:
                continue
            key = preferred.lower()
            if key in seen:
                continue
            seen.add(key)
            lines.append(preferred)

        return lines or ["N/A"]

    # ------------------------------ template rendering ------------------------------
    def _select_template_path(self, state: TaskState, report_name: str) -> Optional[Path]:
        template_path = Path(__file__).resolve().parents[1] / "templates" / "cnvd_template.docx"
        if template_path.exists():
            return template_path

        if self._build_clean_template_from_reports(state, report_name, template_path):
            return template_path
        return None

    def _build_clean_template_from_reports(
        self,
        state: TaskState,
        report_name: str,
        target_path: Path,
    ) -> bool:
        """
        对比“可信代码库报告”和“cnvd报告”同名文档，整理出干净模板（仅结构，不带旧截图/旧内容）。
        """
        try:
            from docx import Document
        except Exception as error:
            self.logger.warning(f"无法创建模板（python-docx 不可用）: {error}")
            return False

        trusted_doc = Path(str(state.input_file))
        if not trusted_doc.exists():
            trusted_doc = Path("可信代码库报告") / f"{report_name}.docx"
        if not trusted_doc.exists():
            trusted_candidates = sorted(Path("可信代码库报告").glob("*.docx")) if Path("可信代码库报告").exists() else []
            trusted_doc = trusted_candidates[0] if trusted_candidates else trusted_doc

        cnvd_doc = Path("cnvd报告") / f"{report_name}.docx"
        if not cnvd_doc.exists():
            cnvd_candidates = sorted(Path("cnvd报告").glob("*.docx")) if Path("cnvd报告").exists() else []
            cnvd_doc = cnvd_candidates[0] if cnvd_candidates else None

        if not trusted_doc.exists() or not cnvd_doc or not Path(cnvd_doc).exists():
            return False

        try:
            trusted_headings = self._extract_heading_order(Path(trusted_doc))
            cnvd_headings = self._extract_heading_order(Path(cnvd_doc))
            line_counts = self._extract_section_text_line_counts(Path(cnvd_doc))
        except Exception as error:
            self.logger.warning(f"模板对比解析失败: {error}")
            return False

        default_order = ["产品介绍", "存在漏洞的代码文件", "漏洞评分", "漏洞原理", "漏洞URL", "复现流程"]
        heading_order = cnvd_headings or default_order
        for heading in default_order:
            if heading not in heading_order:
                heading_order.append(heading)

        target_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("")

        for heading in heading_order:
            doc.add_heading(heading, level=1)
            count = max(1, int(line_counts.get(heading, 1)))
            if heading == "复现流程":
                count = max(count, 4)
            for _ in range(count):
                doc.add_paragraph("")
            if heading in ("产品介绍", "存在漏洞的代码文件", "漏洞评分", "复现流程"):
                doc.add_paragraph("")

        doc.save(target_path)
        self.logger.info(
            f"已基于对比生成模板: {target_path} | trusted_headings={trusted_headings} | cnvd_headings={cnvd_headings}"
        )
        return True

    def _extract_heading_order(self, doc_path: Path) -> List[str]:
        from docx import Document

        doc = Document(doc_path)
        headings: List[str] = []
        seen = set()
        for para in doc.paragraphs:
            style_name = str(getattr(para.style, "name", "") or "")
            text = str(para.text or "").strip()
            if not text or not style_name.startswith("Heading"):
                continue
            key = self._normalize_text_key(text)
            if key in seen:
                continue
            seen.add(key)
            headings.append(text)
        return headings

    def _extract_section_text_line_counts(self, doc_path: Path) -> Dict[str, int]:
        from docx import Document

        doc = Document(doc_path)
        counts: Dict[str, int] = {}
        current_heading = ""
        for para in doc.paragraphs:
            style_name = str(getattr(para.style, "name", "") or "")
            text = str(para.text or "").strip()
            if style_name.startswith("Heading") and text:
                current_heading = text
                counts.setdefault(current_heading, 0)
                continue
            if current_heading and text:
                counts[current_heading] = counts.get(current_heading, 0) + 1
        return counts

    def _fill_report_from_template(self, doc_path: Path, data: Dict[str, Any]) -> None:
        from docx import Document

        doc = Document(doc_path)
        self._clear_all_images(doc)
        source_url = data.get("source_download_url", "") or data.get("repository_url", "") or "N/A"

        section_map = {
            "产品介绍": [
                data.get("product_description", "") or "N/A",
                f"源码下载链接：\t{source_url}",
            ],
            "存在漏洞的代码文件": data.get("vulnerable_files", []) or ["N/A"],
            "漏洞评分": [f"CVSS3.1：{data.get('cvss_score', 'N/A')}"],
            "漏洞原理": self._split_principle_lines(str(data.get("vulnerability_principle", "") or "")),
            "漏洞URL": [data.get("vulnerable_url", "") or "N/A"],
            "复现流程": self._build_reproduction_lines(data),
        }

        for heading, lines in section_map.items():
            self._replace_section_content(doc, heading, lines)

        self._insert_trusted_section_images(doc, data.get("trusted_section_images", {}) or {})
        screenshot = str(data.get("sqlmap_screenshot", "") or "").strip()
        self._insert_screenshot(doc, screenshot)
        doc.save(doc_path)

    def _create_basic_report(self, output_path: Path, data: Dict[str, Any]) -> None:
        from docx import Document

        doc = Document()
        source_url = data.get("source_download_url", "") or data.get("repository_url", "") or "N/A"

        doc.add_paragraph("")
        doc.add_heading("产品介绍", level=1)
        doc.add_paragraph(data.get("product_description", "") or "N/A")
        doc.add_paragraph(f"源码下载链接：\t{source_url}")
        doc.add_paragraph("")

        doc.add_heading("存在漏洞的代码文件", level=1)
        doc.add_paragraph("")
        for line in data.get("vulnerable_files", []) or ["N/A"]:
            doc.add_paragraph(str(line))
        doc.add_paragraph("")

        doc.add_heading("漏洞评分", level=1)
        doc.add_paragraph(f"CVSS3.1：{data.get('cvss_score', 'N/A')}")
        doc.add_paragraph("")

        doc.add_heading("漏洞原理", level=1)
        for line in self._split_principle_lines(str(data.get("vulnerability_principle", "") or "")):
            doc.add_paragraph(line)

        doc.add_heading("漏洞URL", level=1)
        doc.add_paragraph(data.get("vulnerable_url", "") or "N/A")

        doc.add_heading("复现流程", level=1)
        for line in self._build_reproduction_lines(data):
            doc.add_paragraph(line)
        doc.add_paragraph("")
        doc.add_paragraph("")

        self._insert_trusted_section_images(doc, data.get("trusted_section_images", {}) or {})
        screenshot = str(data.get("sqlmap_screenshot", "") or "").strip()
        self._insert_screenshot(doc, screenshot)

        doc.save(output_path)

    def _build_reproduction_lines(self, data: Dict[str, Any]) -> List[str]:
        raw_steps = data.get("reproduction_steps", [])
        lines: List[str] = []
        if isinstance(raw_steps, list):
            lines = [str(item).strip() for item in raw_steps if str(item).strip()]
        elif str(raw_steps or "").strip():
            lines = [str(raw_steps).strip()]
        if not lines:
            lines = ["N/A"]

        # 复现流程必须体现 sqlmap 运行命令（优先展示报告/人工可复用命令）
        sqlmap_command_hint = re.sub(r"\s+", " ", str(data.get("sqlmap_command_hint", "") or "").strip()).strip()
        sqlmap_command_executed = re.sub(
            r"\s+", " ", str(data.get("sqlmap_command_executed", "") or "").strip()
        ).strip()
        sqlmap_command = sqlmap_command_hint or sqlmap_command_executed

        has_sqlmap_line = any("sqlmap " in str(line).lower() for line in lines)
        if sqlmap_command and not has_sqlmap_line:
            lines.append("sqlmap运行指令：")
            lines.append(sqlmap_command)
        elif sqlmap_command:
            has_exact_command = any(sqlmap_command in str(line) for line in lines)
            if not has_exact_command:
                lines.append(f"实际运行命令：{sqlmap_command}")

        if (
            sqlmap_command_hint
            and sqlmap_command_executed
            and sqlmap_command_hint != sqlmap_command_executed
            and not any(sqlmap_command_executed in str(line) for line in lines)
        ):
            lines.append(f"自动化实际执行命令：{sqlmap_command_executed}")

        if not data.get("sqlmap_executed", False):
            fail_reason = str(data.get("sqlmap_failure_reason", "") or "").strip() or "sqlmap 未执行成功"
            lines.append(f"自动复现状态：失败（{fail_reason}）")
            fail_error = str(data.get("sqlmap_error", "") or "").strip()
            if fail_error:
                lines.append(f"失败详情：{fail_error}")
        return [self._pretty_repro_line(line) for line in lines]

    def _pretty_repro_line(self, text: str) -> str:
        line = str(text or "")
        if "sqlmap " in line and "\nsqlmap " not in line:
            line = line.replace("：sqlmap ", "：\nsqlmap ")
            line = line.replace(": sqlmap ", ":\nsqlmap ")
        return line

    def _split_principle_lines(self, text: str) -> List[str]:
        value = str(text or "").replace("\r\n", "\n").strip()
        if not value:
            return ["N/A"]
        lines: List[str] = []
        for block in value.split("\n"):
            current = block.strip()
            if not current:
                continue
            parts = re.split(r"(?<=[。！？])", current)
            for part in parts:
                sentence = part.strip()
                if sentence:
                    lines.append(sentence)
        return lines or ["N/A"]

    def _replace_section_content(self, doc: Any, heading_text: str, lines: List[str]) -> bool:
        heading_idx = self._find_heading_index(doc, heading_text)
        if heading_idx < 0:
            return False

        paragraphs = doc.paragraphs
        next_heading_idx = len(paragraphs)
        for idx in range(heading_idx + 1, len(paragraphs)):
            para = paragraphs[idx]
            style_name = str(getattr(para.style, "name", "") or "")
            if style_name.startswith("Heading") and para.text.strip():
                next_heading_idx = idx
                break

        content_paragraphs = doc.paragraphs[heading_idx + 1 : next_heading_idx]
        normalized_lines = [str(line) for line in (lines or [])] or ["N/A"]

        if content_paragraphs:
            nonblank = [para for para in content_paragraphs if para.text.strip()]
            style = content_paragraphs[0].style
            if nonblank:
                for i, line in enumerate(normalized_lines):
                    if i < len(nonblank):
                        nonblank[i].text = line
                        style = nonblank[i].style
                    else:
                        inserted = self._insert_paragraph_after(nonblank[-1], line, style=style)
                        nonblank.append(inserted)
                if len(nonblank) > len(normalized_lines):
                    for para in nonblank[len(normalized_lines) :]:
                        para.text = ""
            else:
                content_paragraphs[0].text = normalized_lines[0]
                style = content_paragraphs[0].style
                anchor = content_paragraphs[0]
                for line in normalized_lines[1:]:
                    anchor = self._insert_paragraph_after(anchor, line, style=style)
        else:
            heading_para = doc.paragraphs[heading_idx]
            style = heading_para.style
            anchor = self._insert_paragraph_after(heading_para, normalized_lines[0], style=style)
            for line in normalized_lines[1:]:
                anchor = self._insert_paragraph_after(anchor, line, style=style)
        return True

    def _insert_screenshot(self, doc: Any, screenshot_path: str) -> None:
        self._clear_images_in_section(doc, "复现流程")
        path = Path(screenshot_path) if screenshot_path else None
        if not path or not path.exists():
            if screenshot_path:
                self.logger.warning(f"sqlmap 截图不存在，跳过插入: {screenshot_path}")
            return
        max_width = self._get_content_width_emu(doc)
        idx = self._find_heading_index(doc, "复现流程")
        if idx >= 0:
            section_paragraphs = doc.paragraphs[idx + 1 :]
            blank_candidates = [para for para in section_paragraphs if not para.text.strip()]
            if blank_candidates:
                if max_width:
                    blank_candidates[0].add_run().add_picture(str(path), width=max_width)
                else:
                    blank_candidates[0].add_run().add_picture(str(path))
                return
            last_para = doc.paragraphs[min(len(doc.paragraphs) - 1, idx)]
            for para in section_paragraphs:
                last_para = para
            pic_para = self._insert_paragraph_after(last_para, "", style=last_para.style)
            if max_width:
                pic_para.add_run().add_picture(str(path), width=max_width)
            else:
                pic_para.add_run().add_picture(str(path))
            return
        self._append_screenshot_to_doc(doc, str(path))

    def _append_screenshot_to_doc(self, doc: Any, screenshot_path: str) -> None:
        path = Path(screenshot_path)
        if not path.exists():
            return
        max_width = self._get_content_width_emu(doc)
        para = doc.add_paragraph("")
        if max_width:
            para.add_run().add_picture(str(path), width=max_width)
        else:
            para.add_run().add_picture(str(path))

    def _insert_trusted_section_images(self, doc: Any, section_images: Dict[str, Any]) -> None:
        if not isinstance(section_images, dict):
            return

        for heading, paths in section_images.items():
            target_heading = self._map_source_heading_to_target(str(heading))
            if not target_heading or target_heading == "复现流程":
                continue

            valid_paths: List[Path] = []
            if isinstance(paths, list):
                for raw in paths:
                    path = Path(str(raw))
                    if path.exists():
                        valid_paths.append(path)
            if not valid_paths:
                continue

            self._clear_images_in_section(doc, target_heading)
            self._insert_images_into_section(doc, target_heading, valid_paths)

    def _insert_images_into_section(self, doc: Any, heading_text: str, image_paths: List[Path]) -> None:
        heading_idx = self._find_heading_index(doc, heading_text)
        if heading_idx < 0:
            return
        max_width = self._get_content_width_emu(doc)

        paragraphs = doc.paragraphs
        next_heading_idx = len(paragraphs)
        for idx in range(heading_idx + 1, len(paragraphs)):
            para = paragraphs[idx]
            style_name = str(getattr(para.style, "name", "") or "")
            if style_name.startswith("Heading") and para.text.strip():
                next_heading_idx = idx
                break

        section_paragraphs = paragraphs[heading_idx + 1 : next_heading_idx]
        blank_paragraphs = [para for para in section_paragraphs if not para.text.strip()]
        anchor = section_paragraphs[-1] if section_paragraphs else paragraphs[heading_idx]

        for i, image_path in enumerate(image_paths):
            if i < len(blank_paragraphs):
                pic_para = blank_paragraphs[i]
            else:
                pic_para = self._insert_paragraph_after(anchor, "", style=anchor.style)
            if max_width:
                pic_para.add_run().add_picture(str(image_path), width=max_width)
            else:
                pic_para.add_run().add_picture(str(image_path))
            anchor = pic_para
            anchor = self._insert_paragraph_after(anchor, "", style=anchor.style)

    def _clear_all_images(self, doc: Any) -> None:
        for para in doc.paragraphs:
            self._remove_drawings_in_paragraph(para)

    def _clear_images_in_section(self, doc: Any, heading_text: str) -> None:
        start_idx = self._find_heading_index(doc, heading_text)
        if start_idx < 0:
            return
        end_idx = len(doc.paragraphs)
        for idx in range(start_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[idx]
            style_name = str(getattr(para.style, "name", "") or "")
            if style_name.startswith("Heading") and para.text.strip():
                end_idx = idx
                break
        for para in doc.paragraphs[start_idx + 1 : end_idx]:
            self._remove_drawings_in_paragraph(para)

    def _remove_drawings_in_paragraph(self, para: Any) -> None:
        try:
            drawing_nodes = para._p.xpath(".//w:drawing")
            for drawing in drawing_nodes:
                parent = drawing.getparent()
                if parent is not None:
                    parent.remove(drawing)
        except Exception:
            pass

    # ------------------------------ OLE embedding ------------------------------
    def _try_embed_vulnerable_files_as_ole(self, doc_path: Path, file_objects: List[Dict[str, str]]) -> None:
        if not file_objects:
            return

        success, inserted, error_text = self._embed_via_com(doc_path, file_objects)
        if success:
            self.logger.info(f"OLE 嵌入成功（直接 COM），数量: {inserted}")
            return

        self.logger.warning(f"直接 COM 嵌入失败: {error_text}")
        fallback_success = False
        fallback_msg = ""
        if self._looks_like_session_issue(error_text):
            fallback_success, fallback_msg = self._embed_via_schtasks(doc_path, file_objects)
            if fallback_success:
                self.logger.info("OLE 嵌入成功（计划任务交互会话回退）")
                return
            self.logger.warning(f"计划任务回退也失败: {fallback_msg}")

        must_succeed = bool(config.get("output.ole_must_succeed", True))
        if must_succeed:
            detail = error_text if not fallback_msg else f"{error_text} | fallback={fallback_msg}"
            raise RuntimeError(f"OLE 嵌入失败: {detail}")

    def _embed_via_com(self, doc_path: Path, file_objects: List[Dict[str, str]]) -> Tuple[bool, int, str]:
        try:
            import pythoncom  # type: ignore
            import win32com.client  # type: ignore
        except Exception as error:
            return False, 0, f"pywin32 不可用: {error}"

        last_error = ""
        for _ in range(2):
            word = None
            doc = None
            inserted = 0
            try:
                pythoncom.CoInitialize()
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0
                doc = word.Documents.Open(str(doc_path.resolve()))

                self._clear_existing_ole_objects(doc)
                for item in file_objects:
                    if self._insert_single_ole(doc, item):
                        inserted += 1

                doc.Save()
                expected = len(file_objects)
                if inserted >= max(1, expected):
                    return True, inserted, ""
                return False, inserted, f"only inserted {inserted}/{expected} OLE objects"
            except Exception as error:
                last_error = str(error)
                time.sleep(1)
            finally:
                try:
                    if doc is not None:
                        doc.Close(False)
                except Exception:
                    pass
                try:
                    if word is not None:
                        word.Quit()
                except Exception:
                    pass
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
        return False, 0, last_error

    def _clear_existing_ole_objects(self, doc: Any) -> None:
        try:
            count = int(doc.InlineShapes.Count)
            for idx in range(count, 0, -1):
                shape = doc.InlineShapes(idx)
                shape_type = int(getattr(shape, "Type", 0))
                # 1: embedded OLE, 2: linked OLE
                if shape_type in (1, 2):
                    shape.Delete()
        except Exception:
            pass

    def _insert_single_ole(self, doc: Any, item: Dict[str, str]) -> bool:
        source_path = str(item.get("source_path", "") or "").strip()
        if not source_path:
            return False
        file_path = Path(source_path)
        if not file_path.exists():
            self.logger.warning(f"漏洞文件不存在，跳过 OLE: {source_path}")
            return False

        icon_label = str(item.get("display_name", "") or file_path.name)
        original_path = str(item.get("original_path", "") or "").strip()

        target_range = None
        for target_text in [original_path, file_path.name]:
            if target_text:
                target_range = self._find_text_range(doc, target_text)
                if target_range is not None:
                    break
        if target_range is None:
            target_range = self._find_text_range(doc, "存在漏洞的代码文件")
            if target_range is None:
                target_range = doc.Content

        temp_embed_file = self._prepare_embed_file(file_path)
        temp_embed_txt = self._prepare_embed_text_copy(file_path)
        insert_range = target_range.Duplicate
        insert_range.Collapse(0)  # wdCollapseEnd
        insert_range.InsertParagraphAfter()
        insert_range.Collapse(0)

        errors: List[str] = []
        candidate_files = [temp_embed_file]
        if temp_embed_txt:
            candidate_files.append(temp_embed_txt)

        for candidate in candidate_files:
            try:
                if self._insert_ole_attempts(doc, insert_range, candidate, icon_label):
                    tail = doc.Range(insert_range.End, insert_range.End)
                    tail.InsertParagraphAfter()
                    return True
            except Exception as error:
                errors.append(str(error))

        self.logger.warning(f"插入 OLE 失败: {file_path} | {' | '.join(errors)[:800]}")
        return False

    def _insert_ole_attempts(self, doc: Any, insert_range: Any, candidate_file: Path, icon_label: str) -> bool:
        icon_candidates = [
            None,
            r"C:\Windows\System32\packager.dll",
            r"C:\Windows\System32\shell32.dll",
        ]
        attempts = [
            {"ClassType": "Package", "use_selection": False, "use_shape": False},
            {"ClassType": "Package", "use_selection": True, "use_shape": False},
            {"ClassType": None, "use_selection": False, "use_shape": False},
            {"ClassType": "Package", "use_selection": False, "use_shape": True},
        ]

        last_error = ""
        for attempt in attempts:
            for icon in icon_candidates:
                try:
                    kwargs = {
                        "ClassType": attempt["ClassType"],
                        "FileName": str(candidate_file),
                        "LinkToFile": False,
                        "DisplayAsIcon": True,
                        "IconFileName": icon,
                        "IconIndex": 0,
                        "IconLabel": icon_label,
                    }
                    if attempt["use_shape"]:
                        doc.Shapes.AddOLEObject(Anchor=insert_range, **kwargs)
                        return True
                    if attempt["use_selection"]:
                        sel = doc.Application.Selection
                        sel.SetRange(insert_range.Start, insert_range.Start)
                        doc.InlineShapes.AddOLEObject(**kwargs)
                        return True
                    doc.InlineShapes.AddOLEObject(Range=insert_range, **kwargs)
                    return True
                except Exception as error:
                    last_error = str(error)
                    continue
        if last_error:
            raise RuntimeError(last_error)
        return False

    def _prepare_embed_file(self, source_path: Path) -> Path:
        """
        为 OLE 准备一个 ASCII 临时路径副本，规避中文路径/过长路径导致的 Word 命令失败。
        """
        tmp_dir = Path(tempfile.gettempdir()) / "cnvd_ole_files"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        base_name = source_path.name
        safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", base_name)
        target = tmp_dir / safe_name
        if target.exists():
            stamp = int(time.time() * 1000)
            target = tmp_dir / f"{stamp}_{safe_name}"
        try:
            shutil.copy2(source_path, target)
            return target
        except Exception:
            return source_path

    def _prepare_embed_text_copy(self, source_path: Path) -> Optional[Path]:
        tmp_dir = Path(tempfile.gettempdir()) / "cnvd_ole_files"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        base = re.sub(r"[^A-Za-z0-9_.-]", "_", source_path.stem) or "source"
        target = tmp_dir / f"{base}.txt"
        if target.exists():
            target = tmp_dir / f"{base}_{int(time.time() * 1000)}.txt"
        for encoding in ("utf-8", "gbk", "latin1"):
            try:
                text = source_path.read_text(encoding=encoding)
                target.write_text(text, encoding="utf-8")
                return target
            except Exception:
                continue
        return None

    def _looks_like_session_issue(self, error_text: str) -> bool:
        low = str(error_text or "").lower()
        keys = [
            "登录会话不存在",
            "session",
            "0x80070520",
            "0x8000401a",
            "coinitialize",
        ]
        return any(key in low for key in keys)

    def _embed_via_schtasks(self, doc_path: Path, file_objects: List[Dict[str, str]]) -> Tuple[bool, str]:
        helper_script = Path("cnvd_generator/tools/ole_embed_helper.py").resolve()
        if not helper_script.exists():
            return False, f"helper script 不存在: {helper_script}"

        # 计划任务命令行尽量放在 ASCII 路径，避免中文路径被任务计划程序误解析
        tmp_dir = Path(tempfile.gettempdir()) / "cnvd_ole_runtime"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        task_name = f"CNVD_OLE_{int(time.time())}"
        runtime_helper = tmp_dir / "ole_embed_helper_runtime.py"
        try:
            shutil.copy2(helper_script, runtime_helper)
        except Exception:
            runtime_helper = helper_script

        payload_path = tmp_dir / f"{task_name}_payload.json"
        result_path = tmp_dir / f"{task_name}_result.json"
        runner_path = tmp_dir / f"{task_name}.cmd"

        payload = {
            "doc_path": str(doc_path.resolve()),
            "file_objects": file_objects,
        }
        payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        if result_path.exists():
            result_path.unlink()

        run_time = (datetime.now() + timedelta(minutes=1)).strftime("%H:%M")
        command = f"\"{sys.executable}\" \"{runtime_helper}\" --payload \"{payload_path}\" --result \"{result_path}\""
        runner_path.write_text(f"@echo off\r\n{command}\r\n", encoding="utf-8")
        task_tr = f"cmd.exe /c \"{runner_path}\""

        try:
            create = subprocess.run(
                ["schtasks", "/Create", "/F", "/SC", "ONCE", "/TN", task_name, "/TR", task_tr, "/ST", run_time, "/IT"],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=30,
            )
            if create.returncode != 0:
                return False, (create.stderr or create.stdout or "").strip()[:400]

            run = subprocess.run(
                ["schtasks", "/Run", "/TN", task_name],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=30,
            )
            if run.returncode != 0:
                return False, (run.stderr or run.stdout or "").strip()[:400]

            deadline = time.time() + 240
            while time.time() < deadline:
                if result_path.exists():
                    try:
                        result = json.loads(result_path.read_text(encoding="utf-8"))
                    except Exception as error:
                        return False, f"result parse failed: {error}"
                    if bool(result.get("success", False)):
                        return True, f"inserted={result.get('inserted', 0)}"
                    return False, str(result.get("error", "unknown helper error"))
                time.sleep(1)
            return False, "helper result timeout"
        except Exception as error:
            return False, str(error)
        finally:
            try:
                subprocess.run(
                    ["schtasks", "/Delete", "/F", "/TN", task_name],
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    timeout=20,
                )
            except Exception:
                pass
            for p in (payload_path, result_path, runner_path):
                try:
                    if p.exists():
                        p.unlink()
                except Exception:
                    pass

    # ------------------------------ helpers ------------------------------
    def _find_text_range(self, doc: Any, text: str) -> Any:
        rng = doc.Content
        finder = rng.Find
        finder.ClearFormatting()
        finder.Text = str(text)
        finder.Forward = True
        finder.Wrap = 1  # wdFindContinue
        if finder.Execute():
            return rng.Duplicate
        return None

    def _find_heading_index(self, doc: Any, heading_text: str) -> int:
        target = self._normalize_text_key(heading_text)
        for idx, para in enumerate(doc.paragraphs):
            style_name = str(getattr(para.style, "name", "") or "")
            if not style_name.startswith("Heading"):
                continue
            if self._normalize_text_key(para.text) == target:
                return idx
        return -1

    def _normalize_text_key(self, text: str) -> str:
        value = str(text or "")
        value = re.sub(r"\s+", "", value)
        return value.replace("：", ":")

    def _get_content_width_emu(self, doc: Any) -> Optional[int]:
        """获取 Word 正文内容区宽度（EMU），用于约束插图宽度避免版面溢出。"""
        try:
            section = doc.sections[0]
            page_width = int(section.page_width or 0)
            left_margin = int(section.left_margin or 0)
            right_margin = int(section.right_margin or 0)
            content_width = page_width - left_margin - right_margin
            if content_width <= 0:
                return None
            # 预留少量边距，避免贴边
            return int(content_width * 0.95)
        except Exception:
            return None

    def _insert_paragraph_after(self, paragraph: Any, text: str, style: Any = None) -> Any:
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if style is not None:
            new_para.style = style
        if text:
            new_para.add_run(str(text))
        return new_para
